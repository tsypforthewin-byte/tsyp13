# main.py (CV Rewriting Service - FastAPI Version)
from fastapi import FastAPI, HTTPException, Query
from fastapi.responses import StreamingResponse
import tempfile
import os
import json
import logging
from datetime import datetime, timezone
from google.cloud import storage, firestore
from docx import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import ChatOpenAI
import io

app = FastAPI(
    title="CV Rewriter API",
    description="Fetches CV data from Firestore and rewrites it into a formatted DOCX file."
)

# --- Configuration ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY must be set!")

# --- Initialize GCS and Firestore Clients ---
storage_client = storage.Client()
db = firestore.Client(
    project="tsyp-477814",
    database="tsypstore",  # replace with your database ID (e.g., "(default)" or another)
)


GCS_BUCKET_NAME = "tsypbucket"  # Replace with your GCS bucket name

# --- Rewrite CV Function ---
def rewrite_cv(cv_json_data):
    """Rewrites the CV based on the structured JSON data."""
    llm_rewrite = ChatOpenAI(
        model_name="alibaba/tongyi-deepresearch-30b-a3b:free",
        openai_api_key=OPENAI_API_KEY,
        base_url="https://openrouter.ai/api/v1",
        temperature=0.3,
    )

    query_rewrite_cv = """
    You are an expert CV rewriting and formatting assistant.
    You will receive a JSON structure containing extracted CV data from a candidate.

    Your task:
    - Rewrite and format this data into a professional and readable CV in French.
    - Use clear sections (INFORMATIONS PERSONNELLES, PROFIL, FORMATION, EXPÉRIENCE, PROJETS, COMPÉTENCES, CERTIFICATIONS, etc.)
    - Ensure coherence, clarity, and professional tone.
    - If a field is missing or null, simply skip it.
    - Do not invent information.
    - The final output must be a well-structured TEXT (not JSON).
    IMPORTANT : Do not use any markdown syntax (no ##, **, ---, etc.). Use plain text only
    Here is the CV data to rewrite:
    {cv_json}

    Return only the rewritten CV text.
    """

    rewrite_prompt = ChatPromptTemplate.from_template("{prompt}")
    rewrite_chain = (
        {"prompt": RunnablePassthrough()}
        | rewrite_prompt
        | llm_rewrite
        | StrOutputParser()
    )

    cv_json_str = json.dumps(cv_json_data, indent=2, ensure_ascii=False)
    rewritten_cv_text = rewrite_chain.invoke(query_rewrite_cv.format(cv_json=cv_json_str))
    return rewritten_cv_text


# --- Create DOCX Function ---
def create_docx_from_text(cv_text, full_name):
    """Creates a DOCX file from the rewritten CV text and returns a BytesIO object."""
    doc = Document()
    doc.add_heading(f'CV – {full_name}', 0)

    for line in cv_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.isupper() or stripped.endswith(':'):
            doc.add_heading(stripped.rstrip(':'), level=1)
        else:
            doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- API Endpoint ---
@app.get("/rewrite_cv_from_firestore")
async def rewrite_cv_from_firestore_endpoint(user_id: str, cv_doc_id: str):
    """
    Fetches CV data from Firestore using user_id and cv_doc_id, rewrites it, and returns a DOCX file.
    """
    try:
        # --- 1. Fetch CV data from Firestore ---
        doc_ref = db.collection("users").document(user_id).collection("cvs").document(cv_doc_id)
        doc_snapshot = doc_ref.get()

        if not doc_snapshot.exists:
            raise HTTPException(status_code=404, detail=f"CV document '{cv_doc_id}' for user '{user_id}' not found.")

        cv_doc_data = doc_snapshot.to_dict()
        if not cv_doc_data or "cv_data" not in cv_doc_data : 
            raise HTTPException(status_code=500, detail=f"CV document is missing the 'cv_data' field.")

        cv_json = cv_doc_data["cv_data"]
        full_name = cv_json.get('personal_information', {}).get('full_name', 'Unknown')

        # --- 2. Rewrite the CV ---
        rewritten_text = rewrite_cv(cv_json)

        # --- 3. Create the DOCX file in memory ---
        docx_buffer = create_docx_from_text(rewritten_text, full_name)

        # --- 4. Generate a unique filename for GCS ---
        gcs_blob_name = f"rewritten_cvs/{user_id}/{cv_doc_id}_rewritten_{int(datetime.now(timezone.utc).timestamp())}.docx"

        # --- 5. Upload the DOCX buffer to GCS ---
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(gcs_blob_name)
        blob.upload_from_file(
            docx_buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        logging.info(f"✅ Uploaded rewritten CV to GCS: gs://{GCS_BUCKET_NAME}/{gcs_blob_name}")

        # --- 6. Update the original Firestore document with the new GCS path ---
        doc_ref.update({
            "latest_rewritten_gcs_path": f"gs://{GCS_BUCKET_NAME}/{gcs_blob_name}",
            "last_rewritten_at": datetime.now(timezone.utc),
            "rewrite_status": "completed"
        })
        logging.info(f"✅ Updated Firestore document {cv_doc_id} with rewritten file path.")

        # --- 7. Return the DOCX file as a streaming response (download) ---
        docx_buffer.seek(0)
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={full_name}_Rewritten_CV.docx"}
        )

    except HTTPException:
        # Re-raise HTTP exceptions (like 404 Not Found) so FastAPI handles them correctly
        raise
    except Exception as e:
        logging.error(f"Error in rewrite_cv_from_firestore_endpoint: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


# --- Health Check Endpoint ---
@app.get("/")
async def home():
    return {"message": "CV Rewriter API is running!", "framework": "FastAPI"}

@app.get("/health")
async def health():
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))